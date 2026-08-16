"""Génère un document Word avec le schéma de flux détaillé."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def add_code_block(doc, code, font_size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(30, 30, 30)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SAROS — Schéma de flux détaillé")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # =====================================================================
    # 1. VUE GLOBALE
    # =====================================================================
    doc.add_heading("1. Vue globale : Front → Aiguillage → Module", level=1)

    doc.add_paragraph(
        "Lorsqu'un utilisateur saisit une phrase dans l'interface, "
        "voici le parcours complet de sa demande :"
    )

    add_code_block(doc, """\
┌─────────────────────────────────────────────────────────────────────────────┐
│  UTILISATEUR                                                               │
│  "Peux-tu me dire le salaire total de Vincent Dupont sachant qu'il         │
│   gagne 1000 + 200 ?"                                                      │
└────────────────────────────────────┬────────────────────────────────────────┘
                                     │
                                     ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│  INTERFACE FLUTTER / DART                                                   │
│                                                                             │
│  1. Recueille le texte + fichiers éventuels                                │
│  2. Construit le JSON de requête                                           │
│  3. Envoie au service d'aiguillage                                         │
└────────────────────────────────────┬────────────────────────────────────────┘
                                     │  POST JSON
                                     ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│  SERVICE D'AIGUILLAGE (ORCHESTRATION)                                       │
│                                                                             │
│  1. Reçoit la requête                                                      │
│  2. Consulte la table de paramètres                                        │
│  3. Identifie : "anonymisation" → route vers le module                     │
│  4. Appelle POST /anonymisation/detect                                     │
│  5. Renvoie la liste des champs détectés au front                          │
│  6. Reçoit la validation de l'utilisateur                                  │
│  7. Appelle POST /anonymisation/execute                                    │
│  8. Reçoit la réponse finale dé-anonymisée                                 │
│  9. Renvoie la réponse à l'interface                                       │
└────────────────────────────────────┬────────────────────────────────────────┘
                                     │  POST JSON
                                     ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│  MODULE ANONYMISATION (FastAPI)                                             │
│                                                                             │
│  Endpoint 1 : POST /anonymisation/detect                                   │
│  Endpoint 2 : POST /anonymisation/execute                                  │
│                                                                             │
│  ⚠ Tout le traitement sensible se passe ici, côté serveur.                │
│    Les mappings et données anonymisées ne sortent JAMAIS du serveur.        │
└─────────────────────────────────────────────────────────────────────────────┘
""")

    doc.add_page_break()

    # =====================================================================
    # 2. APPEL 1 : DETECT
    # =====================================================================
    doc.add_heading("2. Appel 1 : /anonymisation/detect — Quels fichiers .py sont appelés ?", level=1)

    add_code_block(doc, """\
Orchestration
     │
     │  POST /anonymisation/detect
     │  { "requestId": "...", "message": "...", "files": [...] }
     │
     ▼
┌─ main.py ──────────────────────────────────────────────────────────────────┐
│                                                                             │
│  detect()                                                                  │
│  Rôle : Point d'entrée. Détermine le mode (fichier / PDF / texte)          │
│         et route vers le bon handler interne.                              │
│                                                                             │
│  ┌─ Si fichier Excel/CSV ──────────────────────────────────────────────┐   │
│  │                                                                      │   │
│  │  _detect_file()                                                      │   │
│  │       │                                                              │   │
│  │       ├──► file_handler.py                                          │   │
│  │       │    ├── resolve_file_path()  → trouve le fichier sur disque  │   │
│  │       │    ├── load_file()          → charge en DataFrame pandas    │   │
│  │       │    └── extract_samples()    → prend 3 exemples par colonne  │   │
│  │       │                                                              │   │
│  │       └──► detector.py                                              │   │
│  │            └── detect_sensitive_fields()                             │   │
│  │                 │                                                    │   │
│  │                 └──► Appel LLM local (Ollama)                       │   │
│  │                      via call_local_llm()                           │   │
│  │                      │                                              │   │
│  │                      ├── Envoie colonnes + échantillons             │   │
│  │                      ├── Le LLM analyse et retourne un JSON         │   │
│  │                      └── parse_detection_response()                 │   │
│  └──────────────────────────────────────────────────────────────────────┘   │
│                                                                             │
│  ┌─ Si fichier PDF ────────────────────────────────────────────────────┐   │
│  │                                                                      │   │
│  │  _detect_pdf()                                                       │   │
│  │       │                                                              │   │
│  │       ├──► file_handler.py                                          │   │
│  │       │    ├── resolve_file_path()      → trouve le PDF             │   │
│  │       │    └── extract_text_from_pdf()   → extrait le texte         │   │
│  │       │                                  (via pdfplumber)           │   │
│  │       │                                                              │   │
│  │       └──► detector.py                                              │   │
│  │            └── detect_sensitive_entities()  ← DÉTECTION HYBRIDE     │   │
│  │                                                                      │   │
│  └──────────────────────────────────────────────────────────────────────┘   │
│                                                                             │
│  ┌─ Si texte seul (pas de fichier) ────────────────────────────────────┐   │
│  │                                                                      │   │
│  │  _detect_text()                                                      │   │
│  │       │                                                              │   │
│  │       └──► detector.py                                              │   │
│  │            └── detect_sensitive_entities()  ← DÉTECTION HYBRIDE     │   │
│  │                                                                      │   │
│  └──────────────────────────────────────────────────────────────────────┘   │
│                                                                             │
│  Stocke l'état dans _pending_requests[requestId]                           │
│  Retourne la liste des entités détectées au front                          │
│                                                                             │
└─────────────────────────────────────────────────────────────────────────────┘
""")

    doc.add_page_break()

    # =====================================================================
    # 2b. DÉTECTION HYBRIDE
    # =====================================================================
    doc.add_heading("3. Zoom : La détection hybride (detector.py + pattern_detector.py)", level=1)

    add_code_block(doc, """\
detector.py — detect_sensitive_entities()
     │
     │
     ├──► COUCHE 1 : pattern_detector.py
     │    │
     │    │  detect_by_patterns()
     │    │  Rôle : Détection par expressions régulières (regex)
     │    │         Capture les formats STRUCTURÉS avec fiabilité 100%
     │    │
     │    │  Détecte :
     │    │  ├── IBAN          (BE36 0019 8525 8681)
     │    │  ├── BIC/SWIFT     (GEBABEBB)
     │    │  ├── Email         (vxhaet@gmail.com)
     │    │  ├── Téléphone     (06 01 02 03 04)
     │    │  ├── NIR           (1 85 03 75 108 042 36)
     │    │  ├── TVA belge     (BE 0750.515.724)
     │    │  ├── TVA française (FR 12 345 678 901)
     │    │  └── Communication structurée (+++125/1200/00310+++)
     │    │
     │    └── Retourne : liste de { value, category, label }
     │
     │
     ├──► COUCHE 2 : detector.py (appel LLM local)
     │    │
     │    │  call_local_llm()  → Ollama (qwen2.5:14b-instruct)
     │    │  Rôle : Détection par IA des données NON STRUCTURÉES
     │    │         que les regex ne peuvent pas voir
     │    │
     │    │  Détecte :
     │    │  ├── Noms de personnes    (Vincent Dupont)
     │    │  ├── Noms d'entreprises   (Triple A - Risk Finance Belgium)
     │    │  ├── Adresses postales    (12 rue de la Paix, 75002 Paris)
     │    │  └── Descriptions identifiantes
     │    │
     │    │  Analyse contextuelle :
     │    │  ├── Comprend l'INTENTION de la demande
     │    │  └── NE détecte PAS les montants nécessaires au calcul
     │    │
     │    │  Le prompt utilise rgpd.py pour lister les 27 catégories RGPD
     │    │
     │    └── Retourne : liste de { value, category, strategy, justification }
     │
     │
     └──► FUSION
          │
          │  1. Résultats LLM d'abord (prioritaires, stratégie contextuelle)
          │  2. Résultats regex ensuite (complètent sans doublons)
          │  3. Dédoublonnage par valeur
          │
          └── Retourne : liste fusionnée d'entités détectées
""")

    doc.add_page_break()

    # =====================================================================
    # 3. APPEL 2 : EXECUTE
    # =====================================================================
    doc.add_heading("4. Appel 2 : /anonymisation/execute — Le flux complet côté serveur", level=1)

    doc.add_paragraph(
        "C'est ici que tout se passe. L'utilisateur a validé la liste des champs. "
        "Le front envoie la validation et reçoit directement la réponse finale. "
        "Entre les deux, 3 étapes s'exécutent en interne sur le serveur :"
    )

    add_code_block(doc, """\
Orchestration
     │
     │  POST /anonymisation/execute
     │  { "requestId": "...", "conversationId": "conv-001",
     │    "targetLlm": "claude-sonnet-4-6",
     │    "validatedEntities": [...] }
     │
     ▼
┌─ main.py ──────────────────────────────────────────────────────────────────┐
│                                                                             │
│  execute()                                                                 │
│  Rôle : Orchestre les 3 étapes internes                                    │
│                                                                             │
│  ┌──────────────────────────────────────────────────────────────────────┐   │
│  │ ÉTAPE 1 — ANONYMISATION                          anonymizer.py      │   │
│  │                                                                      │   │
│  │  Anonymizer(existing_mappings=...)                                   │   │
│  │  Rôle : Applique les stratégies validées sur le contenu              │   │
│  │                                                                      │   │
│  │  ├── Charge les mappings existants de la conversation                │   │
│  │  │   (_conversation_mappings[conversationId])                        │   │
│  │  │   → "Vincent Dupont" reste [NOM_1] d'un message à l'autre        │   │
│  │  │                                                                   │   │
│  │  ├── anonymize_text() ou anonymize()                                │   │
│  │  │   ├── Placeholder : "Vincent Dupont" → [NOM_1]                   │   │
│  │  │   │   (tri par longueur décroissante, déduplication)             │   │
│  │  │   └── Chiffrement : "BE36 0019..." → gAAAAABq... (Fernet)       │   │
│  │  │                                                                   │   │
│  │  └── Sauvegarde les mappings mis à jour                             │   │
│  │      _conversation_mappings[conversationId] = mappings               │   │
│  │                                                                      │   │
│  │  Résultat : texte anonymisé                                          │   │
│  │  "Quel est le salaire de [NOM_1] sachant qu il gagne 1000 + 200"    │   │
│  └──────────────────────────────────────────────────────────────────────┘   │
│                          │                                                  │
│                          ▼                                                  │
│  ┌──────────────────────────────────────────────────────────────────────┐   │
│  │ ÉTAPE 2 — APPEL LLM EXTERNE                     llm_router.py      │   │
│  │                                                                      │   │
│  │  _call_llm() → send_to_llm()                                        │   │
│  │  Rôle : Envoie le texte anonymisé au LLM externe                    │   │
│  │                                                                      │   │
│  │  ├── resolve_provider()                                              │   │
│  │  │   "claude-sonnet-4-6" → provider "anthropic"                     │   │
│  │  │   "gpt-4o"            → provider "openai"                        │   │
│  │  │                                                                   │   │
│  │  ├── _send_anthropic() ou _send_openai()                            │   │
│  │  │   Appel HTTP vers l'API du provider                              │   │
│  │  │                                                                   │   │
│  │  │   ┌─────────────────────────────────────────┐                    │   │
│  │  │   │  LLM EXTERNE (Claude / GPT)             │                    │   │
│  │  │   │                                          │                    │   │
│  │  │   │  Reçoit : texte avec [NOM_1], [EMAIL_1] │                    │   │
│  │  │   │  Ne voit JAMAIS les vraies données       │                    │   │
│  │  │   │  Répond : "Le salaire de [NOM_1] est     │                    │   │
│  │  │   │           1200 EUR."                     │                    │   │
│  │  │   └─────────────────────────────────────────┘                    │   │
│  │  │                                                                   │   │
│  │  └── Retourne la réponse brute (avec placeholders)                  │   │
│  └──────────────────────────────────────────────────────────────────────┘   │
│                          │                                                  │
│                          ▼                                                  │
│  ┌──────────────────────────────────────────────────────────────────────┐   │
│  │ ÉTAPE 3 — DÉ-ANONYMISATION                      llm_router.py      │   │
│  │                                                                      │   │
│  │  deanonymize()                                                       │   │
│  │  Rôle : Remet les vraies valeurs dans la réponse du LLM              │   │
│  │                                                                      │   │
│  │  ├── Remplace les placeholders :                                    │   │
│  │  │   [NOM_1] → "Vincent Dupont"                                     │   │
│  │  │   [EMAIL_1] → "vxhaet@gmail.com"                                 │   │
│  │  │                                                                   │   │
│  │  └── Déchiffre les tokens Fernet :                                  │   │
│  │      gAAAAABq... → "BE36 0019 8525 8681"                            │   │
│  │                                                                      │   │
│  │  Résultat : réponse finale lisible                                   │   │
│  │  "Le salaire de Vincent Dupont est 1200 EUR."                        │   │
│  └──────────────────────────────────────────────────────────────────────┘   │
│                                                                             │
│  Retourne la réponse finale dé-anonymisée au front                         │
│  ⚠ Les mappings restent sur le serveur — le front ne les voit JAMAIS      │
│                                                                             │
└─────────────────────────────────────────────────────────────────────────────┘
     │
     ▼
Orchestration → Flutter → Utilisateur
"Le salaire de Vincent Dupont est 1200 EUR."
""")

    doc.add_page_break()

    # =====================================================================
    # 4. CARTE DES FICHIERS
    # =====================================================================
    doc.add_heading("5. Carte des fichiers .py et leur rôle", level=1)

    add_code_block(doc, """\
modules/anonymisation/
│
├── main.py                 CHEF D'ORCHESTRE
│   │                       Expose les 2 endpoints (detect + execute)
│   │                       Route vers les bons handlers
│   │                       Stocke l'état des requêtes et conversations
│   │                       Appelle les autres fichiers dans le bon ordre
│   │
│   ├── appelle ──► file_handler.py      LECTEUR DE FICHIERS
│   │               │                     Charge Excel, CSV, PDF
│   │               │                     Extrait le texte des PDFs
│   │               └── utilise : pandas, openpyxl, pdfplumber
│   │
│   ├── appelle ──► detector.py          CERVEAU DE LA DÉTECTION
│   │               │                     Orchestre la détection hybride
│   │               │                     Construit les prompts pour le LLM
│   │               │                     Fusionne regex + LLM
│   │               │
│   │               ├── appelle ──► pattern_detector.py   DÉTECTION REGEX
│   │               │               │                      Patterns pour IBAN,
│   │               │               │                      TVA, email, tel, NIR
│   │               │               └── utilise : re
│   │               │
│   │               ├── appelle ──► rgpd.py              RÉFÉRENTIEL RGPD
│   │               │               │                      27 catégories
│   │               │               │                      Articles 4, 9, 10
│   │               │               └── get_categories_for_prompt()
│   │               │
│   │               └── appelle ──► Ollama (HTTP)        LLM LOCAL
│   │                               qwen2.5:14b-instruct
│   │                               Détecte noms, entreprises, adresses
│   │
│   ├── appelle ──► anonymizer.py        MOTEUR D'ANONYMISATION
│   │               │                     Applique placeholder ou chiffrement
│   │               │                     Gère les mappings par conversation
│   │               │                     Tri par longueur, déduplication
│   │               └── utilise : cryptography (Fernet)
│   │
│   └── appelle ──► llm_router.py        ROUTEUR LLM EXTERNE
│                   │                     + DÉ-ANONYMISATION
│                   │
│                   ├── send_to_llm()     Envoie au bon provider
│                   │   ├── _send_anthropic()  → API Claude
│                   │   └── _send_openai()     → API GPT
│                   │
│                   └── deanonymize()     Remplace placeholders
│                                          Déchiffre tokens Fernet
│
└── config.py               CONFIGURATION
    │                        Variables d'environnement
    └── SAROS_ANON_*         (Ollama, clés API, stockage)
""")

    doc.add_page_break()

    # =====================================================================
    # 5. RÉSUMÉ SÉQUENTIEL
    # =====================================================================
    doc.add_heading("6. Résumé séquentiel complet", level=1)

    doc.add_paragraph(
        "Voici l'ordre exact de tous les appels de fonctions pour une demande "
        "utilisateur texte, de bout en bout :"
    )

    add_code_block(doc, """\
UTILISATEUR : "Quel est le salaire de Vincent Dupont sachant qu'il gagne 1000+200 ?"

──────────────── APPEL 1 : /detect ────────────────

 1. main.py          → detect()
 2. main.py          → _detect_text()
 3. detector.py      → detect_sensitive_entities()
 4. pattern_detector → detect_by_patterns()           ← regex
 5. detector.py      → build_text_detection_prompt()
 6. rgpd.py          → get_categories_for_prompt()    ← 27 catégories
 7. detector.py      → call_local_llm()               ← appel Ollama
 8. detector.py      → parse_entity_response()
 9. detector.py      → fusion regex + LLM
10. main.py          → stocke dans _pending_requests

→ RETOUR AU FRONT : liste des entités détectées
  [Vincent Dupont → NOM, 17/06/1987 → DATE_NAISS]
  Le front affiche, l'utilisateur valide.

──────────────── APPEL 2 : /execute ────────────────

11. main.py          → execute()
12. main.py          → _execute_text()
13. anonymizer.py    → Anonymizer(existing_mappings=...)  ← charge conversation
14. anonymizer.py    → anonymize_text()                    ← ÉTAPE 1
    "Vincent Dupont" → [NOM_1]
    "1000 + 200" reste en clair (nécessaire au calcul)
15. main.py          → sauvegarde _conversation_mappings
16. main.py          → _call_llm()                         ← ÉTAPE 2
17. llm_router.py    → send_to_llm()
18. llm_router.py    → resolve_provider() → "anthropic"
19. llm_router.py    → _send_anthropic()  → API Claude
    Claude reçoit : "Quel est le salaire de [NOM_1]..."
    Claude répond : "Le salaire de [NOM_1] est 1200 EUR."
20. llm_router.py    → deanonymize()                       ← ÉTAPE 3
    [NOM_1] → "Vincent Dupont"
21. main.py          → retourne la réponse finale

→ RETOUR AU FRONT : "Le salaire de Vincent Dupont est 1200 EUR."
""")

    # Sauvegarde
    output_path = "/Users/vxhaet/projets/Saros/Saros/docs/Saros_Schema_Flux_Detaille.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Document généré : {output_path}")


if __name__ == "__main__":
    main()
