"""Génère le document Word d'architecture globale de Saros."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(30, 30, 30)


def add_arrow_line(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # =====================================================================
    # PAGE DE TITRE
    # =====================================================================
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SAROS")
    run.bold = True
    run.font.size = Pt(40)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Architecture globale de la solution")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Schémas d'architecture et flux d'appels entre composants")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Version 0.1.0 — Août 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # =====================================================================
    # 1. VUE D'ENSEMBLE
    # =====================================================================
    doc.add_heading("1. Vue d'ensemble de Saros", level=1)

    doc.add_paragraph(
        "Saros est une application modulaire qui transforme des demandes utilisateur "
        "en langage naturel en actions concrètes au sein de l'organisation. "
        "L'architecture repose sur trois couches principales :"
    )

    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("Interface utilisateur (Flutter/Dart) : ")
    run.bold = True
    p.add_run("recueille la demande de l'utilisateur sous forme de texte, "
              "avec éventuellement des fichiers joints.")

    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("Service d'orchestration (aiguillage) : ")
    run.bold = True
    p.add_run("collecte la demande et l'oriente vers le bon module fonctionnel "
              "en fonction du type de requête. L'aiguillage est piloté par une "
              "table de paramètres.")

    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("Modules fonctionnels : ")
    run.bold = True
    p.add_run("chaque module possède son propre contexte et domaine. "
              "Le premier module implémenté est le module Anonymisation.")

    # =====================================================================
    # 2. SCHÉMA GLOBAL
    # =====================================================================
    doc.add_heading("2. Schéma d'architecture globale", level=1)

    add_code_block(doc,
        "┌─────────────────────────────────────────────────────────────────────┐\n"
        "│                        UTILISATEUR                                 │\n"
        "│            Saisit une demande en langage naturel                   │\n"
        "│            + fichiers optionnels (Excel, PDF, CSV)                 │\n"
        "└───────────────────────────┬─────────────────────────────────────────┘\n"
        "                            │\n"
        "                            ▼\n"
        "┌─────────────────────────────────────────────────────────────────────┐\n"
        "│                   INTERFACE FLUTTER / DART                         │\n"
        "│                                                                    │\n"
        "│  - Écran de saisie (texte + upload fichiers)                      │\n"
        "│  - Écran de validation (liste des champs sensibles détectés)      │\n"
        "│  - Écran de résultat (réponse finale)                             │\n"
        "└───────────────────────────┬─────────────────────────────────────────┘\n"
        "                            │ JSON (HTTP POST)\n"
        "                            ▼\n"
        "┌─────────────────────────────────────────────────────────────────────┐\n"
        "│              SERVICE D'ORCHESTRATION (AIGUILLAGE)                  │\n"
        "│                                                                    │\n"
        "│  - Reçoit la demande utilisateur                                  │\n"
        "│  - Consulte la table de paramètres                                │\n"
        "│  - Route vers le bon module                                       │\n"
        "│  - Gère le flux d'appels entre les endpoints du module            │\n"
        "│  - Retourne le résultat final à l'interface                       │\n"
        "└──────┬──────────┬──────────┬──────────┬────────────────────────────┘\n"
        "       │          │          │          │\n"
        "       ▼          ▼          ▼          ▼\n"
        "  ┌─────────┐ ┌─────────┐ ┌──────────┐ ┌───────────┐\n"
        "  │ANONYMI- │ │REPOR-   │ │CORRES-   │ │CORRECTION │\n"
        "  │SATION   │ │TING     │ │PONDANCE  │ │DONNÉES    │\n"
        "  │         │ │         │ │          │ │           │\n"
        "  │(actif)  │ │(futur)  │ │(futur)   │ │(futur)    │\n"
        "  └─────────┘ └─────────┘ └──────────┘ └───────────┘"
    )

    doc.add_paragraph(
        "Les échanges entre l'interface, l'orchestration et les modules "
        "se font en JSON via des appels HTTP POST. Chaque module expose "
        "ses propres endpoints REST."
    )

    doc.add_page_break()

    # =====================================================================
    # 3. MODULE ANONYMISATION — ARCHITECTURE INTERNE
    # =====================================================================
    doc.add_heading("3. Module Anonymisation — Architecture interne", level=1)

    doc.add_paragraph(
        "Le module Anonymisation est composé de deux fonctions indépendantes :"
    )

    add_code_block(doc,
        "┌─────────────────────────────────────────────────────────────────────┐\n"
        "│                    MODULE ANONYMISATION                            │\n"
        "│                                                                    │\n"
        "│  ┌─────────────────────────────────────────────────────────────┐   │\n"
        "│  │ FONCTION 1 : ANONYMISATION                                 │   │\n"
        "│  │                                                             │   │\n"
        "│  │  POST /anonymisation/detect      → Détection sensible      │   │\n"
        "│  │  POST /anonymisation/execute      → Anonymisation           │   │\n"
        "│  │  POST /anonymisation/deanonymize  → Dé-anonymisation        │   │\n"
        "│  │                                                             │   │\n"
        "│  │  Composants :                                               │   │\n"
        "│  │  ├── pattern_detector.py  (regex : IBAN, TVA, email...)    │   │\n"
        "│  │  ├── detector.py          (LLM local : noms, entreprises) │   │\n"
        "│  │  ├── anonymizer.py        (placeholder + chiffrement)     │   │\n"
        "│  │  └── rgpd.py              (27 catégories RGPD)            │   │\n"
        "│  └─────────────────────────────────────────────────────────────┘   │\n"
        "│                                                                    │\n"
        "│  ┌─────────────────────────────────────────────────────────────┐   │\n"
        "│  │ FONCTION 2 : ENVOI LLM EXTERNE                             │   │\n"
        "│  │                                                             │   │\n"
        "│  │  POST /llm/send  → Envoie au LLM, retourne réponse brute  │   │\n"
        "│  │                                                             │   │\n"
        "│  │  Composants :                                               │   │\n"
        "│  │  └── llm_router.py  (Anthropic / OpenAI)                  │   │\n"
        "│  │                                                             │   │\n"
        "│  │  ⚠ Aucune connaissance de l'anonymisation                  │   │\n"
        "│  │  → Peut être appelée par n'importe quel module             │   │\n"
        "│  └─────────────────────────────────────────────────────────────┘   │\n"
        "│                                                                    │\n"
        "└─────────────────────────────────────────────────────────────────────┘"
    )

    doc.add_page_break()

    # =====================================================================
    # 4. FLUX D'APPELS DÉTAILLÉ
    # =====================================================================
    doc.add_heading("4. Flux d'appels détaillé — Scénario complet", level=1)

    doc.add_paragraph(
        "Ce schéma montre l'enchaînement complet des appels entre les composants "
        "pour une demande utilisateur nécessitant une anonymisation avant envoi "
        "à un LLM externe."
    )

    add_code_block(doc,
        "UTILISATEUR          FLUTTER            ORCHESTRATION         MODULE ANON.           LLM LOCAL        LLM EXTERNE\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │  saisit demande   │                     │                    │                    │                 │\n"
        "    │  + fichier        │                     │                    │                    │                 │\n"
        "    ├──────────────────►│                     │                    │                    │                 │\n"
        "    │                   │  JSON requête       │                    │                    │                 │\n"
        "    │                   ├────────────────────►│                    │                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │  ① POST /detect    │                    │                 │\n"
        "    │                   │                     ├───────────────────►│                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │                    │  regex detection   │                 │\n"
        "    │                   │                     │                    ├───────────────┐    │                 │\n"
        "    │                   │                     │                    │◄──────────────┘    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │                    │  LLM detection     │                 │\n"
        "    │                   │                     │                    ├───────────────────►│                 │\n"
        "    │                   │                     │                    │◄───────────────────│                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │                    │  fusion résultats  │                 │\n"
        "    │                   │                     │                    ├───────────────┐    │                 │\n"
        "    │                   │                     │                    │◄──────────────┘    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │  liste entités     │                    │                 │\n"
        "    │                   │                     │◄───────────────────│                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │  liste à valider    │                    │                    │                 │\n"
        "    │                   │◄────────────────────│                    │                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │  affiche liste    │                     │                    │                    │                 │\n"
        "    │◄──────────────────│                     │                    │                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │  valide/modifie   │                     │                    │                    │                 │\n"
        "    ├──────────────────►│                     │                    │                    │                 │\n"
        "    │                   │  champs validés     │                    │                    │                 │\n"
        "    │                   ├────────────────────►│                    │                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │  ② POST /execute   │                    │                 │\n"
        "    │                   │                     ├───────────────────►│                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │                    │  anonymise         │                 │\n"
        "    │                   │                     │                    ├───────────────┐    │                 │\n"
        "    │                   │                     │                    │◄──────────────┘    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │  texte anonymisé   │                    │                 │\n"
        "    │                   │                     │  + mappings        │                    │                 │\n"
        "    │                   │                     │◄───────────────────│                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │  ③ POST /llm/send  │                    │                 │\n"
        "    │                   │                     ├───────────────────►│                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │                    │  appel API         │                 │\n"
        "    │                   │                     │                    ├────────────────────────────────────►│\n"
        "    │                   │                     │                    │◄────────────────────────────────────│\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │  réponse brute     │                    │                 │\n"
        "    │                   │                     │  (avec [NOM_1])    │                    │                 │\n"
        "    │                   │                     │◄───────────────────│                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │  ④ POST /deanon.   │                    │                 │\n"
        "    │                   │                     ├───────────────────►│                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │                    │  dé-anonymise      │                 │\n"
        "    │                   │                     │                    ├───────────────┐    │                 │\n"
        "    │                   │                     │                    │◄──────────────┘    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │                     │  réponse finale    │                    │                 │\n"
        "    │                   │                     │◄───────────────────│                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │                   │  réponse finale     │                    │                    │                 │\n"
        "    │                   │◄────────────────────│                    │                    │                 │\n"
        "    │                   │                     │                    │                    │                 │\n"
        "    │  affiche réponse  │                     │                    │                    │                 │\n"
        "    │◄──────────────────│                     │                    │                    │                 │\n"
        "    │                   │                     │                    │                    │                 │"
    )

    doc.add_page_break()

    # =====================================================================
    # 5. DÉTAIL DE CHAQUE APPEL
    # =====================================================================
    doc.add_heading("5. Détail de chaque appel entre composants", level=1)

    # -- Appel 0 : Flutter → Orchestration --
    doc.add_heading("5.0. Interface Flutter → Orchestration", level=2)
    doc.add_paragraph(
        "L'interface Flutter envoie la demande utilisateur au service d'orchestration. "
        "C'est le même format JSON quel que soit le module cible."
    )
    add_code_block(doc,
        'POST https://<orchestration>/api/request\n'
        '\n'
        '{\n'
        '  "requestId": "9f3c8c7e-2d3a-4f2b-8a10-123456789abc",\n'
        '  "userId": "12345",\n'
        '  "conversationId": "conv-789",\n'
        '  "message": "Anonymise les données clients du fichier clients.xlsx",\n'
        '  "files": [\n'
        '    {\n'
        '      "fileId": "file-001",\n'
        '      "name": "clients.xlsx",\n'
        '      "mimeType": "application/vnd.openxmlformats-...",\n'
        '      "size": 152340\n'
        '    }\n'
        '  ],\n'
        '  "context": {\n'
        '    "application": "my_application",\n'
        '    "language": "fr"\n'
        '  }\n'
        '}'
    )
    doc.add_paragraph(
        "L'orchestration consulte sa table de paramètres, identifie que la demande "
        "concerne l'anonymisation, et commence le flux d'appels vers le module."
    )

    # -- Appel 1 --
    doc.add_heading("5.1. Appel ① — POST /anonymisation/detect", level=2)

    doc.add_paragraph("Direction : Orchestration → Module Anonymisation")
    doc.add_paragraph("Objectif : Détecter les données personnelles sensibles dans le contenu.")

    doc.add_paragraph("Requête envoyée :")
    add_code_block(doc,
        'POST https://<module>/anonymisation/detect\n'
        'Content-Type: application/json\n'
        '\n'
        '{\n'
        '  "requestId": "9f3c8c7e-...",\n'
        '  "userId": "12345",\n'
        '  "conversationId": "conv-789",\n'
        '  "message": "Anonymise les données clients du fichier clients.xlsx",\n'
        '  "files": [{...}],\n'
        '  "context": {"application": "my_application", "language": "fr"}\n'
        '}'
    )

    doc.add_paragraph("Traitement interne du module :")
    for text in [
        "Détermine le mode : fichier Excel/CSV, PDF, ou texte seul",
        "Couche 1 (regex) : détecte IBAN, emails, TVA, téléphones, NIR",
        "Couche 2 (LLM local Ollama) : détecte noms, entreprises, adresses",
        "Fusionne les résultats sans doublons",
        "Propose une stratégie par entité (placeholder ou chiffrement)",
        "Stocke l'état en mémoire pour l'étape suivante",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph("Réponse retournée :")
    add_code_block(doc,
        '{\n'
        '  "requestId": "9f3c8c7e-...",\n'
        '  "status": "pending_validation",\n'
        '  "mode": "text",\n'
        '  "detectedEntities": [\n'
        '    {\n'
        '      "value": "Vincent Dupont",\n'
        '      "category": "NOM",\n'
        '      "rgpdArticle": "Article 4",\n'
        '      "recommendedStrategy": "placeholder",\n'
        '      "justification": "Nom identifiant une personne."\n'
        '    },\n'
        '    {\n'
        '      "value": "BE36 0019 8525 8681",\n'
        '      "category": "IBAN",\n'
        '      "rgpdArticle": "Article 4",\n'
        '      "recommendedStrategy": "encryption",\n'
        '      "justification": "Détecté automatiquement (IBAN)."\n'
        '    }\n'
        '  ],\n'
        '  "originalMessage": "..."\n'
        '}'
    )

    doc.add_paragraph(
        "L'orchestration renvoie cette liste à l'interface Flutter. "
        "L'utilisateur la visualise, peut ajouter/supprimer des champs "
        "ou modifier les stratégies, puis valide."
    )

    # -- Appel 2 --
    doc.add_heading("5.2. Appel ② — POST /anonymisation/execute", level=2)

    doc.add_paragraph("Direction : Orchestration → Module Anonymisation")
    doc.add_paragraph("Objectif : Appliquer l'anonymisation avec les choix validés par l'utilisateur.")

    doc.add_paragraph("Requête envoyée :")
    add_code_block(doc,
        'POST https://<module>/anonymisation/execute\n'
        'Content-Type: application/json\n'
        '\n'
        '{\n'
        '  "requestId": "9f3c8c7e-...",\n'
        '  "userId": "12345",\n'
        '  "validatedEntities": [\n'
        '    {"value": "Vincent Dupont", "category": "NOM", "strategy": "placeholder"},\n'
        '    {"value": "BE36 0019 8525 8681", "category": "IBAN", "strategy": "encryption"}\n'
        '  ]\n'
        '}'
    )

    doc.add_paragraph("Traitement interne :")
    for text in [
        "Récupère le contenu original stocké lors du detect",
        "Applique les placeholders : Vincent Dupont → [NOM_1]",
        "Applique le chiffrement Fernet : BE36... → gAAAAABq...",
        "Génère la table de mappings (correspondance placeholder ↔ valeur originale + clé de chiffrement)",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph("Réponse retournée :")
    add_code_block(doc,
        '{\n'
        '  "requestId": "9f3c8c7e-...",\n'
        '  "status": "completed",\n'
        '  "mode": "text",\n'
        '  "anonymizedMessage": "Quel est le salaire de [NOM_1] ...",\n'
        '  "mappings": {\n'
        '    "placeholder_mappings": {\n'
        '      "NOM": {"[NOM_1]": "Vincent Dupont"}\n'
        '    },\n'
        '    "encryption_key": "tAXSB1QHivugy-..."\n'
        '  }\n'
        '}'
    )

    doc.add_paragraph(
        "L'orchestration conserve les mappings en mémoire. "
        "Il en aura besoin à l'étape 4 pour dé-anonymiser la réponse."
    )

    # -- Appel 3 --
    doc.add_heading("5.3. Appel ③ — POST /llm/send", level=2)

    doc.add_paragraph("Direction : Orchestration → Module Anonymisation (Fonction 2)")
    doc.add_paragraph(
        "Objectif : Envoyer le contenu anonymisé à un LLM externe pour obtenir une réponse."
    )

    p = doc.add_paragraph()
    run = p.add_run("Important : ")
    run.bold = True
    p.add_run(
        "cette fonction ne sait rien de l'anonymisation. Elle reçoit du texte, "
        "l'envoie au LLM, et retourne la réponse brute. Elle peut être appelée "
        "par n'importe quel module de Saros."
    )

    doc.add_paragraph("Requête envoyée :")
    add_code_block(doc,
        'POST https://<module>/llm/send\n'
        'Content-Type: application/json\n'
        '\n'
        '{\n'
        '  "requestId": "9f3c8c7e-...",\n'
        '  "userId": "12345",\n'
        '  "content": "Quel est le salaire de [NOM_1] sachant qu il gagne 1000 + 200",\n'
        '  "targetLlm": "claude-sonnet-4-6",\n'
        '  "systemPrompt": "Tu es un assistant financier."  // optionnel\n'
        '}'
    )

    doc.add_paragraph("Traitement interne :")
    for text in [
        'Détermine le provider à partir du nom du modèle (claude-* → Anthropic, gpt-* → OpenAI)',
        "Appelle l'API du provider avec le contenu",
        "Retourne la réponse brute du LLM",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph("Réponse retournée :")
    add_code_block(doc,
        '{\n'
        '  "requestId": "9f3c8c7e-...",\n'
        '  "status": "completed",\n'
        '  "response": "Le salaire total de [NOM_1] est de 1200 EUR.",\n'
        '  "targetLlm": "claude-sonnet-4-6"\n'
        '}'
    )

    doc.add_paragraph(
        "Le LLM externe a répondu avec les placeholders — il n'a jamais vu "
        "les vraies données personnelles."
    )

    # -- Appel 4 --
    doc.add_heading("5.4. Appel ④ — POST /anonymisation/deanonymize", level=2)

    doc.add_paragraph("Direction : Orchestration → Module Anonymisation (Fonction 1)")
    doc.add_paragraph(
        "Objectif : Remettre les vraies valeurs dans la réponse du LLM "
        "en utilisant la table de mappings de l'étape 2."
    )

    doc.add_paragraph("Requête envoyée :")
    add_code_block(doc,
        'POST https://<module>/anonymisation/deanonymize\n'
        'Content-Type: application/json\n'
        '\n'
        '{\n'
        '  "requestId": "9f3c8c7e-...",\n'
        '  "text": "Le salaire total de [NOM_1] est de 1200 EUR.",\n'
        '  "mappings": {\n'
        '    "placeholder_mappings": {\n'
        '      "NOM": {"[NOM_1]": "Vincent Dupont"}\n'
        '    },\n'
        '    "encryption_key": "tAXSB1QHivugy-..."\n'
        '  }\n'
        '}'
    )

    doc.add_paragraph("Traitement interne :")
    for text in [
        "Remplace chaque placeholder par sa valeur originale : [NOM_1] → Vincent Dupont",
        "Détecte les tokens Fernet (commençant par gAAAAA) et les déchiffre avec la clé",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph("Réponse retournée :")
    add_code_block(doc,
        '{\n'
        '  "requestId": "9f3c8c7e-...",\n'
        '  "status": "completed",\n'
        '  "deanonymizedText": "Le salaire total de Vincent Dupont est de 1200 EUR."\n'
        '}'
    )

    doc.add_paragraph(
        "L'orchestration renvoie cette réponse finale à l'interface Flutter "
        "qui l'affiche à l'utilisateur."
    )

    doc.add_page_break()

    # =====================================================================
    # 6. RÉSUMÉ DES FLUX
    # =====================================================================
    doc.add_heading("6. Résumé des flux possibles", level=1)

    doc.add_paragraph(
        "Selon le cas d'usage, l'orchestration n'appelle pas forcément les 4 endpoints. "
        "Voici les combinaisons possibles :"
    )

    add_styled_table(doc,
        ["Cas d'usage", "Endpoints appelés", "Description"],
        [
            [
                "Anonymiser + envoyer au LLM",
                "① detect → ② execute → ③ llm/send → ④ deanonymize",
                "Flux complet : anonymise, envoie au LLM, dé-anonymise la réponse",
            ],
            [
                "Anonymiser uniquement",
                "① detect → ② execute",
                "Anonymise un fichier/texte sans l'envoyer à un LLM",
            ],
            [
                "Envoyer au LLM sans anonymiser",
                "③ llm/send",
                "Appel direct au LLM externe (pas de données sensibles)",
            ],
            [
                "Dé-anonymiser un texte",
                "④ deanonymize",
                "Restaure les données originales à partir d'un texte et de mappings",
            ],
        ],
        col_widths=[4, 6, 7],
    )

    doc.add_page_break()

    # =====================================================================
    # 7. COMPOSANTS EXTERNES
    # =====================================================================
    doc.add_heading("7. Composants externes", level=1)

    add_styled_table(doc,
        ["Composant", "Rôle", "Hébergement", "Configuration"],
        [
            [
                "Ollama",
                "LLM local pour la détection (noms, entreprises, adresses)",
                "Serveur dédié ou local",
                "SAROS_ANON_OLLAMA_BASE_URL\nSAROS_ANON_OLLAMA_MODEL",
            ],
            [
                "API Anthropic",
                "LLM externe (Claude) pour le traitement des demandes",
                "Cloud (api.anthropic.com)",
                "SAROS_ANON_ANTHROPIC_API_KEY",
            ],
            [
                "API OpenAI",
                "LLM externe (GPT) — alternative à Claude",
                "Cloud (api.openai.com)",
                "SAROS_ANON_OPENAI_API_KEY",
            ],
            [
                "Serveur de fichiers",
                "Stockage des fichiers uploadés et anonymisés",
                "Serveur dédié",
                "SAROS_ANON_FILE_STORAGE_PATH",
            ],
        ],
        col_widths=[3, 5, 4, 5],
    )

    # =====================================================================
    # 8. SÉCURITÉ
    # =====================================================================
    doc.add_heading("8. Principes de sécurité", level=1)

    security_items = [
        ("Séparation des responsabilités",
         "La fonction d'envoi LLM n'a aucune connaissance de l'anonymisation. "
         "Les données sensibles ne transitent jamais vers le LLM externe."),
        ("Validation humaine obligatoire",
         "Le module propose une détection mais l'utilisateur doit valider "
         "avant toute anonymisation. Cela garantit qu'aucune donnée sensible "
         "n'est oubliée."),
        ("Chiffrement réversible",
         "Les données chiffrées utilisent Fernet (AES-128-CBC). La clé est "
         "stockée dans la table de mappings, jamais transmise au LLM externe."),
        ("Conformité RGPD",
         "Le référentiel couvre les articles 4 (données personnelles), "
         "9 (données sensibles) et 10 (données pénales) du RGPD."),
        ("Clés API sécurisées",
         "Les clés API sont configurées via variables d'environnement, "
         "jamais stockées dans le code source."),
    ]
    for title_text, desc in security_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title_text + " : ")
        run.bold = True
        p.add_run(desc)

    # -- Sauvegarde --
    output_path = "/Users/vxhaet/projets/Saros/Saros/docs/Saros_Architecture_Globale.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Document généré : {output_path}")


if __name__ == "__main__":
    main()
