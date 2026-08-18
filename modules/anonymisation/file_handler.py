import io
import tempfile

import pdfplumber
import pandas as pd
from pathlib import Path

from .storage import get_file


TABULAR_EXTENSIONS = {".xlsx", ".xls", ".csv"}
TEXT_EXTENSIONS = {".pdf", ".docx", ".txt"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
SUPPORTED_EXTENSIONS = TABULAR_EXTENSIONS | TEXT_EXTENSIONS | IMAGE_EXTENSIONS


def load_file_from_storage(file_id: str, file_name: str) -> pd.DataFrame:
    """Charge un fichier tabulaire depuis MongoDB (ou mémoire)."""
    content = get_file(file_id, file_name)
    if content is None:
        raise FileNotFoundError(
            f"Fichier introuvable dans le stockage : {file_id}/{file_name}"
        )

    suffix = Path(file_name).suffix.lower()
    if suffix not in TABULAR_EXTENSIONS:
        raise ValueError(f"Format non supporté pour le mode tabulaire : {suffix}")

    buf = io.BytesIO(content)
    if suffix in (".xlsx", ".xls"):
        return pd.read_excel(buf)
    return pd.read_csv(buf)


def extract_text_from_storage(file_id: str, file_name: str) -> str:
    """Extrait le texte d'un fichier (PDF, DOCX, TXT, image) depuis le stockage."""
    content = get_file(file_id, file_name)
    if content is None:
        raise FileNotFoundError(
            f"Fichier introuvable dans le stockage : {file_id}/{file_name}"
        )

    suffix = Path(file_name).suffix.lower()

    if suffix == ".pdf":
        return _extract_from_pdf(content)
    elif suffix == ".docx":
        return _extract_from_docx(content)
    elif suffix == ".txt":
        return _extract_from_txt(content)
    elif suffix in IMAGE_EXTENSIONS:
        return _extract_from_image(content)
    else:
        raise ValueError(f"Format non supporté pour l'extraction de texte : {suffix}")


def _extract_from_pdf(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
        tmp.write(content)
        tmp.flush()
        pages_text = []
        with pdfplumber.open(tmp.name) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
    return "\n".join(pages_text)


def _extract_from_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _extract_from_txt(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def _extract_from_image(content: bytes) -> str:
    """Extrait le texte d'une image via OCR (EasyOCR).

    Nécessite un serveur avec suffisamment de RAM (>1GB).
    Désactivé sur Render free (512MB) — actif sur serveur dédié.
    """
    try:
        import easyocr

        with tempfile.NamedTemporaryFile(suffix=".png", delete=True) as tmp:
            tmp.write(content)
            tmp.flush()
            reader = easyocr.Reader(["fr", "en"], gpu=False, verbose=False)
            results = reader.readtext(tmp.name)

        text = " ".join([r[1] for r in results])
        if not text.strip():
            raise ValueError("Aucun texte détecté dans l'image.")
        return text.strip()
    except (MemoryError, OSError):
        raise ValueError(
            "L'OCR d'images nécessite plus de mémoire que disponible sur ce serveur. "
            "Cette fonctionnalité sera active sur le serveur dédié."
        )


def is_text_file(file_name: str) -> bool:
    suffix = Path(file_name).suffix.lower()
    return suffix in TEXT_EXTENSIONS or suffix in IMAGE_EXTENSIONS


def extract_samples(df: pd.DataFrame, n: int = 3) -> dict[str, list[str]]:
    samples = {}
    for col in df.columns:
        non_null = df[col].dropna().head(n)
        samples[col] = [str(v) for v in non_null.tolist()]
    return samples


# ── Génération de fichiers anonymisés ────────────────────────────────


def generate_anonymized_file(
    anonymized_text: str, original_file_name: str
) -> tuple[bytes, str]:
    """Génère un fichier anonymisé dans le même format que l'original.

    Retourne (contenu_bytes, nom_du_fichier_anonymisé).
    """
    suffix = Path(original_file_name).suffix.lower()
    stem = Path(original_file_name).stem
    output_name = f"anonymized_{stem}{suffix}"

    if suffix == ".txt":
        content = _generate_txt(anonymized_text)
    elif suffix == ".docx":
        content = _generate_docx(anonymized_text)
    elif suffix == ".pdf":
        content = _generate_pdf(anonymized_text)
    elif suffix in IMAGE_EXTENSIONS:
        # Pour les images, on génère un TXT (on ne peut pas recréer l'image)
        content = _generate_txt(anonymized_text)
        output_name = f"anonymized_{stem}.txt"
    else:
        content = _generate_txt(anonymized_text)
        output_name = f"anonymized_{stem}.txt"

    return content, output_name


def _generate_txt(text: str) -> bytes:
    return text.encode("utf-8")


def _generate_docx(text: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_pdf(text: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)
    for line in text.split("\n"):
        pdf.cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")
    return pdf.output()
